"""Your boss liked your spreadsheet so much, they would also like a Word document with the same information.

Write a Python program to create  a Word document with the same information. [You can extend the program you wrote
for part 1 or you can write a new program]. It should have a title, and then a list of all of the countries and their
capital cities. Here's an example - the start of your document will look like this:  You can use any styles for the
title and the text."""
# import openpyxl
# import requests
# from openpyxl import Workbook
# import docx
# countries_url = 'https://country-list-1150.azurewebsites.net/api/country'
# country_list_response = requests.get(countries_url).json()


# workbook = openpyxl.load_workbook('counties_and_capitals.xlsx')
# country_workbook = Workbook()
# country_worksheet = country_workbook.active
# countries_url = 'https://country-list-1150.azurewebsites.net/api/country'
# countries_list_response = requests.get(countries_url).json
# document = docx.Document()

# document.add_paragraph('Countries and their Capital Cities', 'Heading 1') for country, capital in
# country_worksheet.items(): document.add_paragraph(f'The capital of {country} is {capital}') document.save(
# 'countries_capitals.doc')
# lines 7 - 24 are my attempt at writing a new code - but after re-reading the Lab
# directions, I saw we could extend previous code instead of writing new code, I thought it said modify,
# so I am going to try that instead
import requests  # must do both
from openpyxl import Workbook  # import library to work with Excel documents in Python -
import docx # all import and from modules need to be at the beginning - I was getting yellow error bars if it was
# after any other code

# this imports just the module we need

#  make request to API to get list of country info
countries_url = 'https://country-list-1150.azurewebsites.net/api/country'  # this is a list of dictionaries, with nested
# dictionaries. usually loop over list, access data by keys
country_list_response = requests.get(countries_url).json()
# descriptive variable names are important
#  create workbook - before loop~

# to have data written to it as loop runs, and we access the country and capital looking things up is expected and
# normal as programmer
# lines 43 - 65 is copy/pasted from Excel Lab - added a few lines of code above and below to
# create, access, and add data needed for word file
country_workbook = Workbook()
country_worksheet = country_workbook.active
country_worksheet.title = 'Countries and Capital Cities'
row = 0
column = 0

document = docx.Document()
document.add_paragraph('Countries and their Capital Cities', 'Title')

#  get the country and capital city from response
for individual_country_dictionary in country_list_response:
    # print(individual_country_dictionary) # huge list, we need to break it down aka access data
    country_name = individual_country_dictionary['name']  # this accesses list of nested dictionaries to give us data
    # based on key 'name'
    # print(country_name) this will print individual country name - good for debugging and checking data
    capital_city = individual_country_dictionary['capitalCity']
    print(f'{country_name:<35} {capital_city}')  # this prints data called from list of dictionaries, format string to
    # use :<35 to add spacing in table
    row = row + 1
    column = column + 1
    country_worksheet.cell(row, 1, country_name)
    country_worksheet.cell(column, 2, capital_city)
    # document.add_paragraph(f'')
    document.add_paragraph(f'The capital of {country_name} is {capital_city}.')
# started by putting docx code outside of loop and making it harder for myself by trying to set up new code after
# loop. realized I needed and was able to add to loop without breaking any of the worksheet code. You can see in
# lines 7-25 I was trying to load the workbook separately but didn't need to try and do that to get data values. Once
# its called, as long as the document.save and workbook.save are listed at the end, the program is able to do both.

country_workbook.save('counties_and_capitals.xlsx')
#  write county and capital city to Workbook
document.save('countries_and_capitals2.docx')
# able to save both while writing code for both as well, as long as the .save is at the end, and related code is in
# loop needed to write to both files.