# may be four or five steps use a combo of everything we've learned goal is to put those ideas together to create a
# longer program leave enough time to finish final - it is a large program and needs attention - get started n o w
# week 1 - start
# week 2 - refining May 1st - pictures
# week 3 - final questions and tweaks May 8th as always, own code and own
# comments
# picking 5 national parks at random, writing to document as a travel guide. Description, weather,
# activities, images (with description and credits - all from api), operating hours, and contact information,
# and website.
# use this url for api to contact - https://national-parks-1150.azurewebsites.net/api/list - first request
# second request - For each park:
# Make a request to the API to get detailed info about the park. You will use a URL in the form
# https://national-parks-1150.azurewebsites.net/api/PARK_CODE but replace the PARK_CODE with the park_code of the
# park chosen. For example,  Acadia National Park has the park_code "ACAD" so the URL for a request for details about
# this park would be to  https://national-parks-1150.azurewebsites.net/api/ACAD.
# Or,  Badlands National Park has the code "BADL" so you will need to use the URL
# https://national-parks-1150.azurewebsites.net/api/BADL Your program
# will need to generate the correct URL to get details about each park you chose.  Hint: these URLs are strings,
# and you can join together strings to make new strings.
# basically, make multiple requests for each park using park_code at end of url api
# 1. make request to get park list
#   2. pick 5 parks at random
#  3. make word doc
#  4. for each park, make another api request to get details from parks, write data from thew response to word doc
#  5. images
#  6. save word doc
#  required comment code as working
# I used this site to help me understand image resizing:
# https://www.geeksforgeeks.org/working-with-images-python-docx-module/

import random  # import random library, imports are usually done at the beginning
import requests  # to get request from API
import docx  # for writing word doc
from docx.shared import Inches

park_word_document = docx.Document()  # writing data to word doc
park_word_document.add_paragraph('National Park Travel Guide', 'Title')
park_list_url = 'https://national-parks-1150.azurewebsites.net/api/list'
# use descriptive names, using more than one
# URL
park_list_response = requests.get(park_list_url).json()
#  accessing a list of dictionaries, which we can access by key
# need random library to pick 5 random parks, we need the park_code
random_park_choice = random.sample(park_list_response, 5)
# accessing random dictionary and obtaining 5 random parks
# we want to get park codes
bullet_point = '\u2022'  # this was used to print activities in a list in Pycharm
# looked up chart for character codes, url: https://altcodeunicode.com/alt-codes-bullet-point-symbols/
for one_park in random_park_choice:
    park_code = one_park['park_code']
    url_for_one_park_details = 'https://national-parks-1150.azurewebsites.net/api/' + park_code
    # urls are strings, this
    # adds random park_code to end or url so, we can access that data
    # print(url_for_one_park_details)
    one_park_details_response = requests.get(url_for_one_park_details).json()
    # we can put requests in the middle of
    # code, does not need to be at the beginning
    park_activities = one_park_details_response['activities']
    activities_list = list(park_activities)
    individual_park_name = one_park_details_response['name']
    individual_park_description = one_park_details_response['description']
    individual_park_location = one_park_details_response['contact_info']['address']
    # attempted to get hours of operation with hours_of_operation = one_park_details_response[operating_hours']-
    # program crashed with "KeyError"
    # attempted to get weather (same thing as operating hours)
    # with park_weather = one_park_details_response['description']['weather']
    park_weather = one_park_details_response['weather_overview']  # got it - key was entered incorrectly
    park_images = one_park_details_response['nps_park_images']
    hours_of_operation = one_park_details_response['park_operating_hours']  # same here as above - wrong key entered
    individual_website_information = one_park_details_response['contact_info']['url']
    # below is code for ease of reading in Python, double-checking word
    print(f'{individual_park_name}')  # if I leave this print it can help me see where the
    # program is at with its process
    # print(f'{individual_park_description}')
    # print(f'{hours_of_operation}')
    # print(f'{park_weather}')
    # print('Activities include:')
    # for activity in park_activities:
    #    print(f' {bullet_point} {activity}')
    # print(f'{individual_park_location}')
    # print(f'{individual_website_information}')
    # test prints to see all data set in variable, I wanted to make sure all the data obtained from APIs were being
    # gathered correctly. Will comment out at another time.
    park_word_document.add_paragraph(individual_park_name, 'Heading 1')
    park_word_document.add_paragraph('Description', 'Heading 2')
    park_word_document.add_paragraph(individual_park_description, 'Normal')
    park_word_document.add_paragraph('Weather', 'Heading 2')
    park_word_document.add_paragraph(park_weather, 'Normal')
    park_word_document.add_paragraph('Activities', 'Heading 2')
    for activity in park_activities:
        park_word_document.add_paragraph(activity, 'List Bullet 2')
        # this part tripped me up a bit, but used the
        # example in the Ppt for word and excel (for car in cars:) and figured it out from there
    park_word_document.add_paragraph('Images', 'Heading 2')
    for image in park_images:
        title = image['title']
        credit = image['credit']
        image_url = image['url']
        nps_images_list = list(park_images)  # this seemed to be an important part for the loop below
        image_response = requests.get(image_url)
        url_parts = image_url.split('/')
        filename = url_parts.pop()
        # print(filename) # also helped with seeing program progress - it also helped me see if the program was writing
        # the .jpg files correctly. This portion really took me some time to figure out
        if nps_images_list == '':
            park_word_document.add_paragraph('No image available', 'Caption')
        for index, url in enumerate(nps_images_list):  # switched the enumerate variable multiple times, needed to
            # tell the program where and how many of the image urls to pull.
            # It kept giving me so many images that would save upwards of 100
            # putting images into nps_images_list solved a lot of issues I was having, pulling images off of park_images
            # variable was not working
            with open(filename, 'wb') as file:
                for chunk in image_response.iter_content():
                    file.write(chunk)
        park_word_document.add_picture(filename, width=Inches(4), height=Inches(4))
        # need to add width and height, images in word document are HUGE. 4x4 seems to be best for formatting
        park_word_document.add_paragraph(f'{title}. {credit}', 'Caption')
        # Below is image code practice - did not go well
        # individual_image_urls = images['url']
        # while park_images == '':
        #    print('No Images available', 'Caption')
        # for index, image_url in enumerate(individual_image_urls):
        # gathered form 'kittens' video, very useful while I was stuck here
        #   image_response = requests.get(individual_image_urls)  # added request in loop to make sure it applied to all
        #    filename = f'nps_image_{index}.jpg'  # generates unique filename for multiple files
        #    with open(filename, 'wb') as file:  # downloads images as files to then write to word doc
        #        for chunk in image_response.iter_content():
        #            file.write(chunk)
        # park_word_document.add_picture()

        # careful with indentation
        # used a format string to add both variables to same line
    park_word_document.add_paragraph('Operating Hours', 'Heading 2')
    park_word_document.add_paragraph(hours_of_operation, 'Normal')
    park_word_document.add_paragraph('Contact Information', 'Heading 2')
    park_word_document.add_paragraph('Address', 'Heading 3')
    park_word_document.add_paragraph(individual_park_location, 'Normal')
    park_word_document.add_paragraph('Website', 'Heading 3')
    park_word_document.add_paragraph(individual_website_information, 'Normal')

park_word_document.save('NationalParkFinalTaylorMc.docx')
# final comments - for the images portion, I decided to use 4 by 4 for spacing. This part took me the longest to figure
# out and troubleshoot. I had certain processes (like the for loop while downloading) had variables (like image response
# and filename data) that would either cause the program to run too long, break, or download a bunch of pictures, or
# overwrite existing ones, causing multiple of the same image to be imported into the Word doc. I also found the
# program would "like" it more if I put filename in the for loop above, AND more specific filenames (building off
# existing filename in url) - the first portion of the program, i.e. writing to the Word doc and getting information for
# description, hours of operation, ect I was able to complete much faster.
# also putting the images into nps_images_list seemed to make it, so I wasn't getting countless pictures to download, as
# well as placing that in the correct loop.
